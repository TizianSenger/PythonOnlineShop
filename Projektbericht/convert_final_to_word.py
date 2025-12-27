#!/usr/bin/env python3
"""
Convert finalProjektbericht.md to Word (.docx) with Professional Formatting
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure paths
BERICHT_DIR = Path(__file__).parent
INPUT_FILE = BERICHT_DIR / "finalProjektbericht.md"
OUTPUT_FILE = BERICHT_DIR / "finalProjektbericht.docx"

def parse_markdown_final(content: str) -> list:
    """Parse markdown and return list of (type, content) tuples"""
    lines = content.split('\n')
    blocks = []
    current_block = []
    current_type = None
    in_code = False
    code_lang = ""
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Check for code blocks
        if line.startswith('```'):
            if in_code:
                blocks.append(('code', '\n'.join(current_block), code_lang))
                current_block = []
                in_code = False
                code_lang = ""
            else:
                if current_block:
                    blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                    current_block = []
                in_code = True
                code_lang = line[3:].strip()
                current_block = []
            i += 1
            continue
        
        if in_code:
            current_block.append(line)
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            blocks.append(('heading1', line[2:].strip()))
            current_type = None
        elif line.startswith('## '):
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            blocks.append(('heading2', line[3:].strip()))
            current_type = None
        elif line.startswith('### '):
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            blocks.append(('heading3', line[4:].strip()))
            current_type = None
        elif line.strip().startswith('|'):
            if current_block and current_type != 'table':
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            current_block.append(line)
            current_type = 'table'
        elif line.strip() == '':
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
                current_type = None
        else:
            if not current_block:
                current_type = 'paragraph'
            current_block.append(line)
        
        i += 1
    
    if current_block:
        blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
    
    return blocks

def add_paragraph(doc: Document, text: str, style: str = 'Normal') -> None:
    """Add paragraph to document"""
    if text.strip():
        p = doc.add_paragraph(text, style=style)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def add_heading(doc: Document, text: str, level: int) -> None:
    """Add heading to document"""
    heading = doc.add_heading(text, level=min(level, 3))
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    heading.paragraph_format.space_after = Pt(6)

def add_code_block(doc: Document, code: str, lang: str = "") -> None:
    """Add code block to document"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    for line in code.split('\n'):
        if line.strip():  # Skip empty lines in code blocks
            run = p.add_run(line + '\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)

def parse_table(table_text: str) -> tuple:
    """Parse markdown table"""
    lines = [l.strip() for l in table_text.split('\n') if l.strip() and l.startswith('|')]
    if len(lines) < 2:
        return None
    
    headers = [h.strip() for h in lines[0].split('|')[1:-1]]
    rows = []
    
    for line in lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells and len(cells) == len(headers):
            rows.append(cells)
    
    return headers, rows if rows else None

def add_table(doc: Document, headers: list, rows: list) -> None:
    """Add table to document with formatting"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Format header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Add rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = cell

def set_page_margins(doc: Document) -> None:
    """Set page margins to 2cm all around"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.787)  # 2cm
        section.bottom_margin = Inches(0.787)
        section.left_margin = Inches(0.787)
        section.right_margin = Inches(0.787)

def add_page_break(doc: Document) -> None:
    """Add page break"""
    doc.add_page_break()

def convert_final_to_docx():
    """Main conversion function for final report"""
    print("🔄 Converting finalProjektbericht.md to Word...")
    
    doc = Document()
    set_page_margins(doc)
    
    # Read markdown file
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse blocks
    blocks = parse_markdown_final(content)
    
    print(f"📊 Found {len(blocks)} content blocks")
    
    # Process blocks
    block_count = 0
    for block_type, block_content, *extra in blocks:
        block_count += 1
        
        if block_type == 'heading1':
            if block_count > 1:  # Page break before each main heading (except first)
                add_page_break(doc)
            add_heading(doc, block_content, 1)
            
        elif block_type == 'heading2':
            add_heading(doc, block_content, 2)
            
        elif block_type == 'heading3':
            add_heading(doc, block_content, 3)
            
        elif block_type == 'code':
            lang = extra[0] if extra else ""
            add_code_block(doc, block_content, lang)
            
        elif block_type == 'table':
            parsed = parse_table(block_content)
            if parsed:
                headers, rows = parsed
                if rows:
                    add_table(doc, headers, rows)
                    
        elif block_type == 'paragraph':
            add_paragraph(doc, block_content)
    
    # Save document
    doc.save(str(OUTPUT_FILE))
    
    # Calculate approximate pages
    approx_pages = len(doc.paragraphs) // 30 + 1
    
    print(f"\n✅ Conversion complete!")
    print(f"📁 Saved to: {OUTPUT_FILE}")
    print(f"📊 Approximate pages: {approx_pages}")
    print(f"📄 File size: {OUTPUT_FILE.stat().st_size / 1024:.2f} KB")

if __name__ == '__main__':
    convert_final_to_docx()
