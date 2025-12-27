#!/usr/bin/env python3
"""
Convert Markdown Projektbericht to Word (.docx)
"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configure paths
BERICHT_DIR = Path(__file__).parent / "Projektbericht"
OUTPUT_FILE = Path(__file__).parent / "Projektbericht" / "PROJEKTBERICHT_Webshop-Python.docx"

# Markdown files in order
MD_FILES = [
    "00_Frontmatter.md",
    "01_Einleitung_und_Anforderungsanalyse.md",
    "02_Zahlungsabwicklung_und_Compliance.md",
    "03_UI_Design_und_Datenmodell.md",
    "04_Technologieentscheidungen.md",
    "05_Architektur_und_Software_Design.md",
    "06_Implementierung_und_MVP.md",
    "07_Testing_und_Qualitaet.md",
    "08_Kritische_Reflexion.md",
    "09_Fazit_und_Ausblick.md",
    "10_Anhang.md",
]

def parse_markdown(content: str) -> list:
    """Parse markdown and return list of (type, content) tuples"""
    lines = content.split('\n')
    blocks = []
    current_block = []
    current_type = None
    in_code = False
    code_lang = ""
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Check for code blocks
        if line.startswith('```'):
            if in_code:
                # End code block
                blocks.append(('code', '\n'.join(current_block), code_lang))
                current_block = []
                in_code = False
                code_lang = ""
            else:
                # Start code block
                in_code = True
                code_lang = line[3:].strip()
                current_block = []
            i += 1
            continue
        
        if in_code:
            current_block.append(line)
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            blocks.append(('heading1', line[2:].strip()))
            current_type = None
        elif line.startswith('## '):
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            blocks.append(('heading2', line[3:].strip()))
            current_type = None
        elif line.startswith('### '):
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            blocks.append(('heading3', line[4:].strip()))
            current_type = None
        elif line.startswith('#### '):
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            blocks.append(('heading4', line[5:].strip()))
            current_type = None
        elif line.strip().startswith('|'):
            # Table row
            if current_block and current_type != 'table':
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
            current_block.append(line)
            current_type = 'table'
        elif line.strip() == '':
            if current_block:
                blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
                current_block = []
                current_type = None
        else:
            if not current_block:
                current_type = 'paragraph'
            current_block.append(line)
        
        i += 1
    
    if current_block:
        blocks.append((current_type or 'paragraph', '\n'.join(current_block)))
    
    return blocks

def add_paragraph(doc: Document, text: str, style: str = 'Normal') -> None:
    """Add paragraph to document"""
    if text.strip():
        p = doc.add_paragraph(text, style=style)

def add_heading(doc: Document, text: str, level: int) -> None:
    """Add heading to document"""
    if level == 1:
        doc.add_heading(text, level=1)
    elif level == 2:
        doc.add_heading(text, level=2)
    elif level == 3:
        doc.add_heading(text, level=3)
    else:
        doc.add_heading(text, level=min(level, 4))

def add_code_block(doc: Document, code: str, lang: str = "") -> None:
    """Add code block to document"""
    p = doc.add_paragraph(code, style='Normal')
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.left_indent = Inches(0.5)

def parse_table(table_text: str) -> list:
    """Parse markdown table"""
    lines = [l.strip() for l in table_text.split('\n') if l.strip() and l.startswith('|')]
    if len(lines) < 2:
        return None
    
    # Parse header
    headers = [h.strip() for h in lines[0].split('|')[1:-1]]
    
    # Skip separator line
    # Parse rows
    rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    return headers, rows if rows else None

def add_table(doc: Document, headers: list, rows: list) -> None:
    """Add table to document"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Add rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = cell

def convert_markdown_to_docx():
    """Main conversion function"""
    print("🔄 Converting Markdown Projektbericht to Word...")
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('PROJEKTBERICHT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('WEBSHOP-PYTHON: Konzeption und Umsetzung eines Onlineshops')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    # Process all markdown files
    for md_file in MD_FILES:
        file_path = BERICHT_DIR / md_file
        if not file_path.exists():
            print(f"⚠️  Skipping {md_file} (not found)")
            continue
        
        print(f"📄 Processing {md_file}...")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse blocks
        blocks = parse_markdown(content)
        
        for block_type, block_content, *extra in blocks:
            if block_type == 'heading1':
                doc.add_page_break()
                add_heading(doc, block_content, 1)
            elif block_type == 'heading2':
                add_heading(doc, block_content, 2)
            elif block_type == 'heading3':
                add_heading(doc, block_content, 3)
            elif block_type == 'heading4':
                add_heading(doc, block_content, 4)
            elif block_type == 'code':
                lang = extra[0] if extra else ""
                add_code_block(doc, block_content, lang)
            elif block_type == 'table':
                parsed = parse_table(block_content)
                if parsed:
                    headers, rows = parsed
                    if rows:
                        add_table(doc, headers, rows)
            elif block_type == 'paragraph':
                add_paragraph(doc, block_content)
    
    # Save document
    doc.save(str(OUTPUT_FILE))
    print(f"\n✅ Conversion complete!")
    print(f"📁 Saved to: {OUTPUT_FILE}")
    print(f"📊 File size: {OUTPUT_FILE.stat().st_size / 1024 / 1024:.2f} MB")

if __name__ == '__main__':
    convert_markdown_to_docx()
